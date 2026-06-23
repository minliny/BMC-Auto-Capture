from __future__ import annotations

import csv
import json
import zipfile
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.out.acceptance_docx import generate_acceptance_artifacts


def _write_template(path: Path) -> None:
    doc = Document()
    doc.add_heading("测试用例及测试记录", level=1)
    doc.add_heading("管理功能测试", level=2)
    doc.add_heading("计算节点部件信息查询测试", level=3)
    table = doc.add_table(rows=8, cols=2)
    labels = ["用例编号", "测试目的", "测试组网", "预置条件", "测试步骤", "预期结果", "测试结果", "备注"]
    values = ["4.2.4", "目的", "NA", "", "", "", "", ""]
    for idx, label in enumerate(labels):
        table.rows[idx].cells[0].text = label
        table.rows[idx].cells[1].text = values[idx]

    doc.add_heading("测试结果分析", level=1)
    summary = doc.add_table(rows=2, cols=4)
    for idx, value in enumerate(["测试类别", "用例编号", "用例名称", "测试结果 （pass/fail/NT）"]):
        summary.rows[0].cells[idx].text = value
    for idx, value in enumerate(["管理功能测试", "4.2.4", "计算节点部件信息查询测试", ""]):
        summary.rows[1].cells[idx].text = value
    doc.save(path)


def _append_case_to_template(path: Path, case_id: str, title: str, category: str = "管理功能测试") -> None:
    doc = Document(path)
    doc.add_heading(category, level=2)
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=8, cols=2)
    labels = ["用例编号", "测试目的", "测试组网", "预置条件", "测试步骤", "预期结果", "测试结果", "备注"]
    values = [case_id, "目的", "NA", "", "", "", "", ""]
    for idx, label in enumerate(labels):
        table.rows[idx].cells[0].text = label
        table.rows[idx].cells[1].text = values[idx]
    doc.save(path)


def _write_png(path: Path, color: tuple[int, int, int]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    Image.new("RGB", (640, 360), color).save(path)


def _picture_paragraphs(doc: Document):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph._p.xpath(".//w:drawing"):
                        yield paragraph


def _summary_verdict_cells(doc: Document, verdict: str):
    for table in doc.tables:
        if len(table.rows) < 2 or len(table.columns) < 4:
            continue
        header = [" ".join(cell.text.split()) for cell in table.rows[0].cells[:4]]
        if header[:3] != ["测试类别", "用例编号", "用例名称"]:
            continue
        for row in table.rows[1:]:
            if len(row.cells) >= 4 and " ".join(row.cells[3].text.split()) == verdict:
                yield row.cells[3]


def test_generate_acceptance_docx_and_evidence_zip(tmp_path: Path):
    run_output = tmp_path / "output" / "20260623_103000"
    template = tmp_path / "template.docx"
    _write_template(template)

    cpu_png = run_output / "4.2.4_计算节点部件信息查询测试-CPU" / "A3" / "10.0.0.1-计算节点部件信息查询测试-CPU.png"
    npu_png = run_output / "4.2.4_计算节点部件信息查询测试-NPU" / "A3" / "10.0.0.1-计算节点部件信息查询测试-NPU.png"
    ssh_png = run_output / "4.1.10_NPU驱动和固件安装测试" / "A3" / "10.0.0.2-NPU驱动和固件安装测试.png"
    ssh_txt = run_output / "4.1.10_NPU驱动和固件安装测试" / "A3" / "10.0.0.2-NPU驱动和固件安装测试.txt"
    _write_png(cpu_png, (200, 20, 20))
    _write_png(npu_png, (20, 200, 20))
    _write_png(ssh_png, (20, 20, 200))
    ssh_txt.parent.mkdir(parents=True, exist_ok=True)
    ssh_txt.write_text("npu-smi info\nOK\n", encoding="utf-8")

    csv_path = run_output / "final_result.csv"
    csv_path.parent.mkdir(parents=True, exist_ok=True)
    with csv_path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["任务序号", "任务名称", "任务类型", "设备分类", "设备名称", "执行状态", "最终结论", "执行失败原因", "截图路径", "文本路径"])
        writer.writerow(["4.2.4", "计算节点部件信息查询测试-CPU", "BMC", "A3", "A3-01", "EXEC_SUCCESS", "PASS", "", str(cpu_png), ""])
        writer.writerow(["4.2.4", "计算节点部件信息查询测试-NPU", "BMC", "A3", "A3-01", "EXEC_SUCCESS", "PASS", "", str(npu_png), ""])
        writer.writerow(["4.1.10", "NPU驱动和固件安装测试", "SSH", "A3", "A3-01", "EXEC_SUCCESS", "PASS", "", str(ssh_png), str(ssh_txt)])

    result = generate_acceptance_artifacts(
        run_output=run_output,
        template_path=template,
        output_dir=run_output,
    )

    assert result.docx_path.exists()
    assert result.evidence_zip_path.exists()
    assert result.report_path.exists()
    assert result.matched_cases == 1
    assert result.packaged_files == 4

    filled = Document(result.docx_path)
    text = "\n".join(paragraph.text for paragraph in filled.paragraphs)
    for table in filled.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    assert "测试结果：PASS" in text
    assert "证据1：计算节点部件信息查询测试-CPU" in text
    assert "证据2：计算节点部件信息查询测试-NPU" in text
    summary_verdict_cells = list(_summary_verdict_cells(filled, "PASS"))
    assert summary_verdict_cells
    for cell in summary_verdict_cells:
        paragraph = cell.paragraphs[0]
        assert paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT
        assert paragraph.paragraph_format.left_indent.pt == 0
        assert paragraph.paragraph_format.first_line_indent.pt == 0
    picture_paragraphs = list(_picture_paragraphs(filled))
    assert picture_paragraphs
    for paragraph in picture_paragraphs:
        assert paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT
        assert paragraph.paragraph_format.left_indent.pt == 0
        assert paragraph.paragraph_format.first_line_indent.pt == 0
    with zipfile.ZipFile(result.docx_path) as zf:
        import xml.etree.ElementTree as ET

        xml = zf.read("word/document.xml")
    root = ET.fromstring(xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    drawing_cells = [tc for tc in root.findall(".//w:tc", ns) if tc.find(".//w:drawing", ns) is not None]
    assert drawing_cells
    for tc in drawing_cells:
        left = tc.find("./w:tcPr/w:tcMar/w:left", ns)
        assert left is not None
        assert left.attrib["{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w"] == "0"
        grid_span = tc.find("./w:tcPr/w:gridSpan", ns)
        assert grid_span is not None
        assert int(grid_span.attrib["{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"]) >= 2
    nowrap_cells = [
        tc for tc in root.findall(".//w:tc", ns)
        if "".join(text.text or "" for text in tc.findall(".//w:t", ns)) == "PASS"
    ]
    assert nowrap_cells
    for tc in nowrap_cells:
        assert tc.find("./w:tcPr/w:noWrap", ns) is not None

    with zipfile.ZipFile(result.evidence_zip_path) as zf:
        names = sorted(zf.namelist())
    assert "4.2.4_计算节点部件信息查询测试-CPU/A3/10.0.0.1-计算节点部件信息查询测试-CPU.png" in names
    assert "4.2.4_计算节点部件信息查询测试-NPU/A3/10.0.0.1-计算节点部件信息查询测试-NPU.png" in names
    assert "4.1.10_NPU驱动和固件安装测试/A3/10.0.0.2-NPU驱动和固件安装测试.png" in names
    assert "4.1.10_NPU驱动和固件安装测试/A3/10.0.0.2-NPU驱动和固件安装测试.txt" in names
    assert all(not name.startswith("output/") for name in names)
    assert all("20260623_103000" not in name for name in names)

    report = json.loads(result.report_path.read_text(encoding="utf-8"))
    assert report["cases"][0]["matched_tasks"] == [
        "计算节点部件信息查询测试-CPU",
        "计算节点部件信息查询测试-NPU",
    ]
    assert report["unmatched_case_details"] == []
    assert [item["task_name"] for item in report["unmatched_result_tasks"]] == ["NPU驱动和固件安装测试"]
    assert report["unmatched_result_tasks"][0]["task_sequence"] == "4.1.10"
    assert report["unmatched_result_tasks"][0]["task_type"] == "SSH"


def test_generate_acceptance_docx_filters_selected_evidence_dirs(tmp_path: Path):
    run_output = tmp_path / "output" / "20260623_103000"
    template = tmp_path / "template.docx"
    _write_template(template)

    cpu_dir = run_output / "4.2.4_计算节点部件信息查询测试-CPU" / "A3"
    npu_dir = run_output / "4.2.4_计算节点部件信息查询测试-NPU" / "A3"
    cpu_png = cpu_dir / "10.0.0.1-计算节点部件信息查询测试-CPU.png"
    npu_png = npu_dir / "10.0.0.1-计算节点部件信息查询测试-NPU.png"
    _write_png(cpu_png, (200, 20, 20))
    _write_png(npu_png, (20, 200, 20))

    csv_path = run_output / "final_result.csv"
    csv_path.parent.mkdir(parents=True, exist_ok=True)
    with csv_path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["任务序号", "任务名称", "任务类型", "设备分类", "设备名称", "执行状态", "最终结论", "执行失败原因", "截图路径", "文本路径"])
        writer.writerow(["4.2.4", "计算节点部件信息查询测试-CPU", "BMC", "A3", "A3-01", "EXEC_SUCCESS", "PASS", "", str(cpu_png), ""])
        writer.writerow(["4.2.4", "计算节点部件信息查询测试-NPU", "BMC", "A3", "A3-01", "EXEC_SUCCESS", "PASS", "", str(npu_png), ""])

    result = generate_acceptance_artifacts(
        evidence_dirs=[cpu_dir],
        template_path=template,
        output_dir=run_output,
    )

    report = json.loads(result.report_path.read_text(encoding="utf-8"))
    assert report["run_output"] == str(run_output.resolve())
    assert report["cases"][0]["matched_tasks"] == ["计算节点部件信息查询测试-CPU"]
    assert report["unmatched_case_details"] == []
    assert report["unmatched_result_tasks"] == []
    assert result.packaged_files == 1

    with zipfile.ZipFile(result.evidence_zip_path) as zf:
        names = sorted(zf.namelist())
    assert names == [
        "4.2.4_计算节点部件信息查询测试-CPU/A3/10.0.0.1-计算节点部件信息查询测试-CPU.png"
    ]


def test_generate_acceptance_docx_remaps_migrated_result_paths(tmp_path: Path):
    old_root = tmp_path / "old_output" / "20260623_103000"
    new_root = tmp_path / "migrated" / "20260623_103000"
    template = tmp_path / "template.docx"
    _write_template(template)

    old_cpu_png = old_root / "4.2.4.计算节点部件信息查询测试-CPU" / "A3" / "10.0.0.1-计算节点部件信息查询测试-CPU.png"
    old_npu_png = old_root / "4.2.4.计算节点部件信息查询测试-NPU" / "A3" / "10.0.0.1-计算节点部件信息查询测试-NPU.png"
    new_cpu_dir = new_root / "4.2.4_计算节点部件信息查询测试-CPU" / "A3"
    new_npu_dir = new_root / "4.2.4_计算节点部件信息查询测试-NPU" / "A3"
    new_cpu_png = new_cpu_dir / old_cpu_png.name
    new_npu_png = new_npu_dir / old_npu_png.name
    _write_png(new_cpu_png, (200, 20, 20))
    _write_png(new_npu_png, (20, 200, 20))

    csv_path = new_root / "final_result.csv"
    csv_path.parent.mkdir(parents=True, exist_ok=True)
    with csv_path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["任务序号", "任务名称", "任务类型", "设备分类", "设备名称", "执行状态", "最终结论", "执行失败原因", "截图路径", "文本路径"])
        writer.writerow(["4.2.4", "计算节点部件信息查询测试-CPU", "BMC", "A3", "A3-01", "EXEC_SUCCESS", "PASS", "", str(old_cpu_png), ""])
        writer.writerow(["4.2.4", "计算节点部件信息查询测试-NPU", "BMC", "A3", "A3-01", "EXEC_SUCCESS", "PASS", "", str(old_npu_png), ""])

    result = generate_acceptance_artifacts(
        evidence_dirs=[new_cpu_dir],
        template_path=template,
        output_dir=new_root,
    )

    assert result.matched_cases == 1
    assert result.packaged_files == 1
    assert result.missing_files == 0

    report = json.loads(result.report_path.read_text(encoding="utf-8"))
    assert report["cases"][0]["matched_tasks"] == ["计算节点部件信息查询测试-CPU"]
    assert report["missing_file_details"] == []
    assert report["unmatched_case_details"] == []
    assert report["unmatched_result_tasks"] == []

    with zipfile.ZipFile(result.evidence_zip_path) as zf:
        names = sorted(zf.namelist())
    assert names == [
        "4.2.4_计算节点部件信息查询测试-CPU/A3/10.0.0.1-计算节点部件信息查询测试-CPU.png"
    ]


def test_generate_acceptance_docx_synthesizes_rows_from_legacy_dot_dirs_without_result_csv(tmp_path: Path):
    run_output = tmp_path / "output" / "manual"
    template = tmp_path / "template.docx"
    _write_template(template)

    cpu_dir = run_output / "4.2.4.计算节点部件信息查询测试-CPU" / "A3"
    cpu_png = cpu_dir / "10.0.0.1-计算节点部件信息查询测试-CPU.png"
    _write_png(cpu_png, (200, 20, 20))

    result = generate_acceptance_artifacts(
        evidence_dirs=[cpu_dir],
        template_path=template,
        output_dir=run_output,
    )

    assert result.matched_cases == 1
    assert result.packaged_files == 1
    report = json.loads(result.report_path.read_text(encoding="utf-8"))
    assert report["cases"][0]["matched_tasks"] == ["计算节点部件信息查询测试-CPU"]
    assert report["unmatched_case_details"] == []
    assert report["unmatched_result_tasks"] == []


def test_generate_acceptance_docx_matches_underscore_sequence_dirs_without_result_csv(tmp_path: Path):
    run_output = tmp_path / "Pod3-4 常规测试用例截图"
    template = tmp_path / "template.docx"
    _write_template(template)
    _append_case_to_template(template, "4.1.7", "计算节点上电测试", "基本功能测试")
    _append_case_to_template(template, "4.1.8", "RAID配置测试", "基本功能测试")

    evidence_files = [
        run_output / "0.0.0_A3 BMC首页截图" / "A3" / "10.0.0.1-A3 BMC首页截图.png",
        run_output / "4.1.7_计算节点上电测试-CPU" / "A3" / "10.0.0.1-计算节点上电测试-CPU.png",
        run_output / "4.1.7_计算节点上电测试-内存" / "A3" / "10.0.0.1-计算节点上电测试-内存.png",
        run_output / "4.1.8_RAID配置测试" / "A3" / "10.0.0.1-RAID配置测试.png",
        run_output / "4.2.4_计算节点部件信息查询测试-CPU" / "A3" / "10.0.0.1-计算节点部件信息查询测试-CPU.png",
        run_output / "4.2.4_计算节点部件信息查询测试-NPU" / "A3" / "10.0.0.1-计算节点部件信息查询测试-NPU.png",
    ]
    for index, path in enumerate(evidence_files, start=1):
        _write_png(path, (20 * index, 30, 40))

    result = generate_acceptance_artifacts(
        evidence_dirs=[run_output],
        template_path=template,
        output_dir=run_output,
    )

    assert result.matched_cases == 3
    assert result.unmatched_cases == 0

    report = json.loads(result.report_path.read_text(encoding="utf-8"))
    matched_by_title = {case["title"]: case["matched_tasks"] for case in report["cases"]}
    assert matched_by_title["计算节点上电测试"] == [
        "计算节点上电测试-CPU",
        "计算节点上电测试-内存",
    ]
    assert matched_by_title["RAID配置测试"] == ["RAID配置测试"]
    assert matched_by_title["计算节点部件信息查询测试"] == [
        "计算节点部件信息查询测试-CPU",
        "计算节点部件信息查询测试-NPU",
    ]
    assert report["unmatched_case_details"] == []
    assert [item["task_name"] for item in report["unmatched_result_tasks"]] == ["A3 BMC首页截图"]
    assert report["unmatched_result_tasks"][0]["task_sequence"] == "0.0.0"

    with zipfile.ZipFile(result.evidence_zip_path) as zf:
        names = sorted(zf.namelist())
    assert "0.0.0_A3 BMC首页截图/A3/10.0.0.1-A3 BMC首页截图.png" in names
    assert "4.1.7_计算节点上电测试-CPU/A3/10.0.0.1-计算节点上电测试-CPU.png" in names
    assert "4.2.4_计算节点部件信息查询测试-NPU/A3/10.0.0.1-计算节点部件信息查询测试-NPU.png" in names


def test_generate_acceptance_docx_reports_unmatched_template_cases(tmp_path: Path):
    run_output = tmp_path / "output" / "20260623_103000"
    template = tmp_path / "template.docx"
    _write_template(template)
    _append_case_to_template(template, "4.9.9", "未执行测试")

    cpu_png = run_output / "4.2.4_计算节点部件信息查询测试-CPU" / "A3" / "10.0.0.1-计算节点部件信息查询测试-CPU.png"
    _write_png(cpu_png, (200, 20, 20))

    csv_path = run_output / "final_result.csv"
    csv_path.parent.mkdir(parents=True, exist_ok=True)
    with csv_path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["任务序号", "任务名称", "任务类型", "设备分类", "设备名称", "执行状态", "最终结论", "执行失败原因", "截图路径", "文本路径"])
        writer.writerow(["4.2.4", "计算节点部件信息查询测试-CPU", "BMC", "A3", "A3-01", "EXEC_SUCCESS", "PASS", "", str(cpu_png), ""])

    result = generate_acceptance_artifacts(
        run_output=run_output,
        template_path=template,
        output_dir=run_output,
    )

    assert result.matched_cases == 1
    assert result.unmatched_cases == 1

    report = json.loads(result.report_path.read_text(encoding="utf-8"))
    assert report["unmatched_case_details"] == [
        {
            "case_id": "4.9.9",
            "category": "管理功能测试",
            "title": "未执行测试",
        }
    ]
    assert report["unmatched_result_tasks"] == []
    assert [item["status"] for item in report["cases"]] == ["MATCHED", "NO_MATCH"]
