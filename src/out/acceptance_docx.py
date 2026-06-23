"""Fill acceptance DOCX reports from run results and package evidence files."""

from __future__ import annotations

import csv
import json
import logging
import os
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from ..utils.path_safety import safe_filename

logger = logging.getLogger("bmc_auto_capture.acceptance_docx")


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
TASK_DIR_PATTERN = re.compile(r"^(\d+(?:\.\d+)+)[._](.+)$")
DEFAULT_TEMPLATE_RELATIVE_PATH = Path(
    "templates/acceptance/Atlas 900 A3 SuperPoD 超节点 验收测试指南 03.docx"
)


@dataclass
class ResultRow:
    row_index: int
    task_sequence: str
    task_name: str
    task_type: str
    device_group: str
    device_name: str
    execution_status: str
    final_verdict: str
    failure_reason: str
    output_dir: Path | None = None
    screenshot_paths: list[Path] = field(default_factory=list)
    txt_paths: list[Path] = field(default_factory=list)


@dataclass
class CaseBlock:
    case_id: str
    category: str
    title: str
    result_cell: object


@dataclass
class AcceptanceOutput:
    docx_path: Path
    evidence_zip_path: Path
    report_path: Path
    matched_cases: int
    unmatched_cases: int
    packaged_files: int
    missing_files: int


def default_template_path(app_dir: str | Path | None = None) -> Path:
    """Return the bundled acceptance template path."""
    if app_dir:
        return Path(app_dir) / DEFAULT_TEMPLATE_RELATIVE_PATH
    return Path(__file__).resolve().parents[2] / DEFAULT_TEMPLATE_RELATIVE_PATH


def generate_acceptance_artifacts(
    *,
    run_output: str | Path | None = None,
    evidence_dirs: Iterable[str | Path] | None = None,
    template_path: str | Path | None = None,
    output_dir: str | Path | None = None,
    app_dir: str | Path | None = None,
) -> AcceptanceOutput:
    """Create a filled DOCX, evidence ZIP, and fill report from a run directory."""
    selected_dirs = _normalize_evidence_dirs(evidence_dirs)
    run_root = _resolve_run_root(run_output, selected_dirs)

    template = Path(template_path).resolve() if template_path else default_template_path(app_dir).resolve()
    if not template.exists():
        raise FileNotFoundError(f"Acceptance DOCX template not found: {template}")

    dest_dir = Path(output_dir).resolve() if output_dir else run_root
    dest_dir.mkdir(parents=True, exist_ok=True)

    stem = template.stem
    docx_path = dest_dir / f"{stem}_filled.docx"
    evidence_zip_path = dest_dir / f"{stem}_evidence.zip"
    report_path = dest_dir / f"{stem}_fill_report.json"

    results = _read_rows_for_export(run_root, selected_dirs)
    cases, doc = _load_cases(template)
    match_report = _fill_document(doc, cases, results)
    doc.save(docx_path)

    zip_report = package_evidence(run_root, results, evidence_zip_path)
    report = {
        "run_output": str(run_root),
        "template": str(template),
        "evidence_dirs": [str(path) for path in selected_dirs],
        "docx": str(docx_path),
        "evidence_zip": str(evidence_zip_path),
        "matched_cases": match_report["matched_cases"],
        "unmatched_cases": match_report["unmatched_cases"],
        "unmatched_case_details": match_report["unmatched_case_details"],
        "unmatched_result_tasks": match_report["unmatched_result_tasks"],
        "cases": match_report["cases"],
        "packaged_files": zip_report["packaged_files"],
        "missing_files": zip_report["missing_files"],
        "missing_file_details": zip_report["missing_file_details"],
    }
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    return AcceptanceOutput(
        docx_path=docx_path,
        evidence_zip_path=evidence_zip_path,
        report_path=report_path,
        matched_cases=int(match_report["matched_cases"]),
        unmatched_cases=int(match_report["unmatched_cases"]),
        packaged_files=int(zip_report["packaged_files"]),
        missing_files=int(zip_report["missing_files"]),
    )


def read_result_rows(run_root: Path) -> list[ResultRow]:
    csv_path = _find_result_csv(run_root)
    rows: list[ResultRow] = []
    with csv_path.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for index, raw in enumerate(reader, start=2):
            row = ResultRow(
                row_index=index,
                task_sequence=_get(raw, "任务序号", "TaskSequence", "task_sequence"),
                task_name=_get(raw, "任务名称", "TaskName", "task_name"),
                task_type=_get(raw, "任务类型", "TaskType", "task_type").upper(),
                device_group=_get(raw, "设备分类", "设备分组", "DeviceGroup", "device_group"),
                device_name=_get(raw, "设备名称", "DeviceName", "device_name"),
                execution_status=_get(raw, "执行状态", "execution_status"),
                final_verdict=_get(raw, "最终结论", "final_verdict"),
                failure_reason=_get(raw, "执行失败原因", "execution_failure_reason"),
            )
            output_dir = _get(raw, "输出目录", "output_dir")
            if output_dir:
                row.output_dir = _resolve_paths(run_root, [output_dir])[0]
            row.screenshot_paths = _resolve_paths(run_root, _split_paths(_get(raw, "截图路径", "screenshots")))
            row.txt_paths = _resolve_paths(run_root, _split_paths(_get(raw, "文本路径", "txt_file")))
            rows.append(row)
    return rows


def _read_rows_for_export(run_root: Path, selected_dirs: list[Path]) -> list[ResultRow]:
    if _has_result_csv(run_root):
        rows = read_result_rows(run_root)
        if selected_dirs:
            rows = _remap_rows_to_selected_dirs(rows, selected_dirs)
            return _filter_rows_by_dirs(rows, selected_dirs)
        return rows
    if selected_dirs:
        return _synthesize_rows_from_dirs(run_root, selected_dirs)
    raise FileNotFoundError(f"Neither final_result.csv nor result.csv exists under {run_root}")


def _normalize_evidence_dirs(evidence_dirs: Iterable[str | Path] | None) -> list[Path]:
    if not evidence_dirs:
        return []
    dirs: list[Path] = []
    for raw in evidence_dirs:
        if raw is None:
            continue
        text = str(raw).strip().strip('"')
        if not text:
            continue
        path = Path(text).resolve()
        if not path.is_dir():
            raise FileNotFoundError(f"Evidence directory not found: {path}")
        dirs.append(path)
    return dirs


def _resolve_run_root(run_output: str | Path | None, selected_dirs: list[Path]) -> Path:
    if run_output:
        run_root = Path(run_output).resolve()
        if not run_root.is_dir():
            raise FileNotFoundError(f"Run output directory not found: {run_root}")
        return run_root
    for directory in selected_dirs:
        root = _nearest_result_root(directory)
        if root is not None:
            return root
    for directory in selected_dirs:
        root = _nearest_task_parent(directory)
        if root is not None:
            return root
    if selected_dirs:
        return _common_existing_root(selected_dirs)
    raise ValueError("run_output or evidence_dirs is required for acceptance DOCX export")


def _nearest_result_root(path: Path) -> Path | None:
    current = path.resolve()
    candidates = [current, *current.parents]
    for candidate in candidates:
        if _has_result_csv(candidate):
            return candidate
    return None


def _nearest_task_parent(path: Path) -> Path | None:
    current = path.resolve()
    for candidate in [current, *current.parents]:
        if _parse_task_dir_name(candidate.name):
            return candidate.parent
    return None


def _has_result_csv(path: Path) -> bool:
    return any((path / name).exists() for name in ("final_result.csv", "result.csv"))


def _common_existing_root(paths: list[Path]) -> Path:
    try:
        common = Path(os.path.commonpath([str(path) for path in paths])).resolve()
    except ValueError:
        return paths[0].resolve()
    return common if common.is_dir() else common.parent


def _filter_rows_by_dirs(rows: list[ResultRow], selected_dirs: list[Path]) -> list[ResultRow]:
    if not selected_dirs:
        return rows
    return [row for row in rows if _row_matches_selected_dirs(row, selected_dirs)]


def _remap_rows_to_selected_dirs(
    rows: list[ResultRow],
    selected_dirs: list[Path],
) -> list[ResultRow]:
    files = _collect_selected_evidence_files(selected_dirs)
    for row in rows:
        row.screenshot_paths = [
            _remap_file_to_selected_dirs(path, row, selected_dirs, files)
            for path in row.screenshot_paths
        ]
        row.txt_paths = [
            _remap_file_to_selected_dirs(path, row, selected_dirs, files)
            for path in row.txt_paths
        ]
        row.output_dir = _remap_output_dir_to_selected_dirs(row, selected_dirs, files)
    return rows


def _collect_selected_evidence_files(selected_dirs: list[Path]) -> list[Path]:
    files: list[Path] = []
    seen: set[Path] = set()
    for directory in selected_dirs:
        for path in sorted(directory.rglob("*")):
            if not path.is_file():
                continue
            if path.suffix.lower() not in IMAGE_EXTENSIONS and path.suffix.lower() != ".txt":
                continue
            resolved = path.resolve()
            if resolved not in seen:
                seen.add(resolved)
                files.append(resolved)
    return files


def _remap_file_to_selected_dirs(
    path: Path,
    row: ResultRow,
    selected_dirs: list[Path],
    files: list[Path],
) -> Path:
    old_path = path.resolve()
    if any(_path_within(old_path, selected_dir) for selected_dir in selected_dirs):
        return old_path

    suffixes = _candidate_suffixes(old_path, row)
    for suffix in suffixes:
        for selected_dir in selected_dirs:
            candidate = _candidate_under_selected_dir(selected_dir, suffix)
            if candidate and candidate.exists() and candidate.is_file():
                return candidate.resolve()

    matches = [
        candidate for candidate in files
        if candidate.name == old_path.name and _evidence_file_matches_row(candidate, row)
    ]
    if len(matches) == 1:
        return matches[0]

    matches = [candidate for candidate in files if candidate.name == old_path.name]
    if len(matches) == 1:
        return matches[0]

    return old_path


def _remap_output_dir_to_selected_dirs(
    row: ResultRow,
    selected_dirs: list[Path],
    files: list[Path],
) -> Path | None:
    evidence_paths = [
        path for path in [*row.screenshot_paths, *row.txt_paths]
        if any(_path_within(path, selected_dir) for selected_dir in selected_dirs)
    ]
    if evidence_paths:
        return evidence_paths[0].parent

    output_dir = row.output_dir.resolve() if row.output_dir else None
    if output_dir and any(_path_within(output_dir, selected_dir) for selected_dir in selected_dirs):
        return output_dir

    matching_files = [path for path in files if _evidence_file_matches_row(path, row)]
    if matching_files:
        return matching_files[0].parent
    return output_dir


def _candidate_suffixes(path: Path, row: ResultRow) -> list[Path]:
    suffixes: list[Path] = []
    parts = list(path.parts)
    task_index = _task_part_index(parts, row)
    if task_index is not None:
        suffixes.append(Path(*parts[task_index:]))
        if task_index + 1 < len(parts):
            suffixes.append(Path(*parts[task_index + 1:]))
        if task_index + 2 < len(parts):
            suffixes.append(Path(*parts[task_index + 2:]))
    if row.device_group:
        for index, part in enumerate(parts):
            if part == row.device_group and index + 1 < len(parts):
                suffixes.append(Path(*parts[index:]))
                suffixes.append(Path(*parts[index + 1:]))
    suffixes.append(Path(path.name))

    unique: list[Path] = []
    seen: set[str] = set()
    for suffix in suffixes:
        key = suffix.as_posix()
        if key not in seen:
            seen.add(key)
            unique.append(suffix)
    return unique


def _task_part_index(parts: list[str], row: ResultRow) -> int | None:
    sequence = row.task_sequence.strip()
    task_name = row.task_name.strip()
    for index, part in enumerate(parts):
        parsed = _parse_task_dir_name(part)
        if parsed:
            part_sequence, part_task_name = parsed
            sequence_matches = not sequence or part_sequence == sequence
            task_matches = not task_name or _task_names_related(part_task_name, task_name)
            if sequence_matches and task_matches:
                return index
        if sequence and task_name and part in {f"{sequence}.{task_name}", f"{sequence}_{task_name}"}:
            return index
        if sequence and _path_part_has_sequence(part, sequence):
            return index
        if task_name and part == task_name:
            return index
    for index, part in enumerate(parts):
        if task_name and task_name in part:
            return index
    return None


def _candidate_under_selected_dir(selected_dir: Path, suffix: Path) -> Path | None:
    suffix_parts = suffix.parts
    if not suffix_parts:
        return None
    selected_parts = selected_dir.resolve().parts
    drop = 0
    max_drop = min(len(suffix_parts), len(selected_parts))
    while drop < max_drop and selected_parts[-drop - 1] == suffix_parts[drop]:
        drop += 1
    remaining = suffix_parts[drop:]
    if not remaining:
        return selected_dir.resolve()
    return (selected_dir / Path(*remaining)).resolve()


def _evidence_file_matches_row(path: Path, row: ResultRow) -> bool:
    parts = path.resolve().parts
    joined = "/".join(parts)
    if row.task_name and row.task_name not in joined and row.task_name not in path.name:
        return False
    if row.task_sequence and not any(_path_part_has_sequence(part, row.task_sequence) for part in parts):
        return False
    if row.device_group and row.device_group not in parts:
        return False
    return True


def _row_matches_selected_dirs(row: ResultRow, selected_dirs: list[Path]) -> bool:
    paths = [path for path in [row.output_dir, *row.screenshot_paths, *row.txt_paths] if path is not None]
    return any(
        _path_within(path, selected_dir) or _path_within(selected_dir, path)
        for path in paths
        for selected_dir in selected_dirs
    )


def _path_within(path: Path, parent: Path) -> bool:
    try:
        path.resolve().relative_to(parent.resolve())
        return True
    except ValueError:
        return False


def _synthesize_rows_from_dirs(run_root: Path, selected_dirs: list[Path]) -> list[ResultRow]:
    rows_by_key: dict[tuple[str, str, str], ResultRow] = {}
    row_index = 2
    for selected_dir in selected_dirs:
        for path in sorted(selected_dir.rglob("*")):
            if not path.is_file():
                continue
            suffix = path.suffix.lower()
            if suffix not in IMAGE_EXTENSIONS and suffix != ".txt":
                continue
            task_sequence, task_name, device_group = _infer_task_from_path(path, run_root)
            key = (task_sequence, task_name, device_group)
            row = rows_by_key.get(key)
            if row is None:
                row = ResultRow(
                    row_index=row_index,
                    task_sequence=task_sequence,
                    task_name=task_name,
                    task_type="SSH" if suffix == ".txt" else "BMC",
                    device_group=device_group,
                    device_name="",
                    execution_status="EXEC_SUCCESS",
                    final_verdict="PASS",
                    failure_reason="",
                    output_dir=path.parent,
                )
                rows_by_key[key] = row
                row_index += 1
            if suffix in IMAGE_EXTENSIONS:
                row.screenshot_paths.append(path.resolve())
            elif suffix == ".txt":
                row.txt_paths.append(path.resolve())
                row.task_type = "SSH"
    return list(rows_by_key.values())


def _infer_task_from_path(path: Path, run_root: Path) -> tuple[str, str, str]:
    try:
        parts = path.resolve().relative_to(run_root.resolve()).parts
    except ValueError:
        parts = path.resolve().parts

    for index, part in enumerate(parts):
        parsed = _parse_task_dir_name(part)
        if parsed:
            device_group = parts[index + 1] if index + 1 < len(parts) - 1 else ""
            return parsed[0], parsed[1], device_group

    parent = path.parent
    return "", parent.name, ""


def _parse_task_dir_name(name: str) -> tuple[str, str] | None:
    match = TASK_DIR_PATTERN.match(name)
    if not match:
        return None
    return match.group(1), match.group(2).strip()


def _path_part_has_sequence(part: str, sequence: str) -> bool:
    parsed = _parse_task_dir_name(part)
    if parsed:
        return parsed[0] == sequence
    return part.startswith(sequence + ".") or part.startswith(sequence + "_")


def _task_names_related(left: str, right: str) -> bool:
    if left == right:
        return True
    return left.startswith(right + "-") or right.startswith(left + "-")


def package_evidence(run_root: Path, rows: Iterable[ResultRow], evidence_zip_path: Path) -> dict:
    packaged = 0
    missing: list[dict] = []
    seen: set[str] = set()

    with zipfile.ZipFile(evidence_zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for row in rows:
            for path in row.screenshot_paths:
                if path.suffix.lower() not in IMAGE_EXTENSIONS:
                    continue
                if not path.exists():
                    missing.append(_missing_detail(row, path))
                    continue
                arcname = _unique_arcname(_arcname_for(run_root, path), seen)
                zf.write(path, arcname)
                packaged += 1

            if not _is_ssh_or_telnet(row.task_type):
                continue
            for path in row.txt_paths:
                if path.suffix.lower() != ".txt":
                    continue
                if not path.exists():
                    missing.append(_missing_detail(row, path))
                    continue
                arcname = _unique_arcname(_arcname_for(run_root, path), seen)
                zf.write(path, arcname)
                packaged += 1

    return {
        "packaged_files": packaged,
        "missing_files": len(missing),
        "missing_file_details": missing,
    }


def _find_result_csv(run_root: Path) -> Path:
    for name in ("final_result.csv", "result.csv"):
        candidate = run_root / name
        if candidate.exists():
            return candidate
    raise FileNotFoundError(f"Neither final_result.csv nor result.csv exists under {run_root}")


def _get(row: dict, *names: str) -> str:
    for name in names:
        value = row.get(name)
        if value is not None:
            return str(value).strip()
    return ""


def _split_paths(raw: str) -> list[str]:
    return [part.strip() for part in raw.replace("\n", ";").split(";") if part.strip()]


def _resolve_paths(run_root: Path, paths: Iterable[str]) -> list[Path]:
    resolved: list[Path] = []
    for raw in paths:
        p = Path(raw)
        if p.is_absolute():
            resolved.append(p)
            continue
        candidates = [
            run_root / p,
            run_root.parent / p,
            run_root.parent.parent / p,
            Path.cwd() / p,
        ]
        for candidate in candidates:
            if candidate.exists():
                resolved.append(candidate.resolve())
                break
        else:
            resolved.append((run_root / p).resolve())
    return resolved


def _load_cases(template: Path):
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise RuntimeError("python-docx is required to export acceptance DOCX reports") from exc

    doc = Document(template)
    last_h2 = ""
    last_h3 = ""
    cases: list[CaseBlock] = []

    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = Paragraph(child, doc)
            text = _norm(paragraph.text)
            style_name = paragraph.style.name if paragraph.style else ""
            if text and style_name == "Heading 2":
                last_h2 = text
            elif text and style_name == "Heading 3":
                last_h3 = text
        elif child.tag.endswith("}tbl"):
            table = Table(child, doc)
            case_id = _case_id_from_table(table)
            if not case_id:
                continue
            result_cell = _result_cell_from_table(table)
            if result_cell is None:
                continue
            cases.append(CaseBlock(case_id=case_id, category=last_h2, title=last_h3, result_cell=result_cell))
    return cases, doc


def _case_id_from_table(table) -> str:
    if len(table.rows) < 2 or len(table.columns) < 2:
        return ""
    if _norm(table.rows[0].cells[0].text) != "用例编号":
        return ""
    return _norm(table.rows[0].cells[1].text)


def _result_cell_from_table(table):
    for row in table.rows:
        if len(row.cells) >= 2 and _norm(row.cells[0].text) == "测试结果":
            return row.cells[0].merge(row.cells[-1])
    return None


def _fill_document(doc, cases: list[CaseBlock], rows: list[ResultRow]) -> dict:
    cases_report: list[dict] = []
    used_row_ids: set[int] = set()
    matched = 0
    unmatched = 0

    for case in cases:
        case_rows = _match_rows(case.title, rows)
        if not case_rows:
            unmatched += 1
            cases_report.append({
                "case_id": case.case_id,
                "category": case.category,
                "title": case.title,
                "status": "NO_MATCH",
                "matched_tasks": [],
            })
            continue

        matched += 1
        used_row_ids.update(id(row) for row in case_rows)
        selected = _one_image_per_task(case_rows)
        verdict = _case_verdict(case_rows)
        _write_case_result(case.result_cell, verdict, selected)
        cases_report.append({
            "case_id": case.case_id,
            "category": case.category,
            "title": case.title,
            "status": "MATCHED",
            "verdict": verdict,
            "matched_tasks": [row.task_name for row in case_rows],
            "inserted_tasks": [row.task_name for row in selected],
        })

    _update_summary_table(doc, cases_report)
    unmatched_case_details = [
        {
            "case_id": item["case_id"],
            "category": item.get("category", ""),
            "title": item["title"],
        }
        for item in cases_report
        if item.get("status") == "NO_MATCH"
    ]
    return {
        "matched_cases": matched,
        "unmatched_cases": unmatched,
        "unmatched_case_details": unmatched_case_details,
        "unmatched_result_tasks": [
            _result_row_detail(row)
            for row in rows
            if id(row) not in used_row_ids
        ],
        "cases": cases_report,
    }


def _result_row_detail(row: ResultRow) -> dict:
    return {
        "row_index": row.row_index,
        "task_sequence": row.task_sequence,
        "task_name": row.task_name,
        "task_type": row.task_type,
        "device_group": row.device_group,
        "device_name": row.device_name,
        "execution_status": row.execution_status,
        "final_verdict": row.final_verdict,
        "output_dir": str(row.output_dir) if row.output_dir else "",
        "screenshot_paths": [str(path) for path in row.screenshot_paths],
        "txt_paths": [str(path) for path in row.txt_paths],
    }


def _match_rows(case_title: str, rows: list[ResultRow]) -> list[ResultRow]:
    exact = [row for row in rows if row.task_name == case_title]
    prefixed = [
        row for row in rows
        if row.task_name.startswith(case_title + "-") and row not in exact
    ]
    return sorted(exact + prefixed, key=lambda row: (row.task_name, row.device_group, row.device_name, row.row_index))


def _one_image_per_task(rows: list[ResultRow]) -> list[ResultRow]:
    selected_by_task: dict[str, ResultRow] = {}
    for row in rows:
        if row.task_name not in selected_by_task:
            selected_by_task[row.task_name] = row
            continue
        current = selected_by_task[row.task_name]
        if not _has_existing_image(current) and _has_existing_image(row):
            selected_by_task[row.task_name] = row
    return list(selected_by_task.values())


def _has_existing_image(row: ResultRow) -> bool:
    return any(path.exists() and path.suffix.lower() in IMAGE_EXTENSIONS for path in row.screenshot_paths)


def _is_ssh_or_telnet(task_type: str) -> bool:
    value = (task_type or "").upper()
    return "SSH" in value or "TELNET" in value


def _case_verdict(rows: list[ResultRow]) -> str:
    values = [(row.final_verdict or row.execution_status or "").upper() for row in rows]
    if any(value in {"FAIL", "EXEC_FAILED", "EXEC_ERROR", "EXEC_TIMEOUT"} for value in values):
        return "FAIL"
    if any(value in {"WARN", "WARNING", "EXEC_PARTIAL"} for value in values):
        return "WARN"
    if values and all(value in {"PASS", "EXEC_SUCCESS", "RULE_DISABLED", ""} for value in values):
        return "PASS"
    if any(value in {"SKIPPED", "BLOCKED"} or value.startswith("EXEC_SKIPPED") for value in values):
        return "NT"
    return values[0] if values else "NT"


def _write_case_result(cell, verdict: str, rows: list[ResultRow]) -> None:
    from docx.shared import Inches

    cell.text = ""
    _set_cell_left_margin(cell, 0)
    first = cell.paragraphs[0]
    _left_align_paragraph(first)
    first.add_run(f"测试结果：{verdict}").bold = True

    for index, row in enumerate(rows, start=1):
        label = cell.add_paragraph()
        _left_align_paragraph(label)
        label.add_run(f"证据{index}：{row.task_name}").bold = True
        device = " / ".join(part for part in (row.device_group, row.device_name) if part)
        if device:
            label.add_run(f"（{device}）")

        image = next((path for path in row.screenshot_paths if path.exists() and path.suffix.lower() in IMAGE_EXTENSIONS), None)
        if image is None:
            missing = cell.add_paragraph("未找到可插入的截图。")
            _left_align_paragraph(missing)
            continue
        picture_paragraph = cell.add_paragraph()
        _left_align_paragraph(picture_paragraph)
        picture_paragraph.add_run().add_picture(str(image), width=Inches(5.1))

        reason = row.failure_reason.strip()
        if reason:
            reason_paragraph = cell.add_paragraph(f"说明：{reason}")
            _left_align_paragraph(reason_paragraph)


def _left_align_paragraph(paragraph) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Pt(0)
    paragraph_format.first_line_indent = Pt(0)


def _set_cell_left_margin(cell, dxa: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    left = tc_mar.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        tc_mar.append(left)
    left.set(qn("w:w"), str(dxa))
    left.set(qn("w:type"), "dxa")


def _update_summary_table(doc, cases_report: list[dict]) -> None:
    verdict_by_title = {
        item["title"]: item.get("verdict", "")
        for item in cases_report
        if item.get("status") == "MATCHED"
    }
    for table in doc.tables:
        if len(table.rows) < 2 or len(table.columns) < 4:
            continue
        header = [_norm(cell.text) for cell in table.rows[0].cells[:4]]
        if header[:3] != ["测试类别", "用例编号", "用例名称"]:
            continue
        for row in table.rows[1:]:
            if len(row.cells) < 4:
                continue
            title = _norm(row.cells[2].text)
            verdict = verdict_by_title.get(title)
            if verdict:
                _write_summary_verdict_cell(row.cells[3], verdict)


def _write_summary_verdict_cell(cell, verdict: str) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_left_margin(cell, 120)
    _set_cell_no_wrap(cell)
    paragraph = cell.paragraphs[0]
    _apply_paragraph_style(paragraph, "TableText")
    _left_align_paragraph(paragraph)
    paragraph.add_run(verdict)


def _apply_paragraph_style(paragraph, style_name: str) -> None:
    for style in paragraph.part.document.styles:
        if style.name == style_name or style.style_id == style_name:
            paragraph.style = style
            return


def _set_cell_no_wrap(cell) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if no_wrap is None:
        no_wrap = OxmlElement("w:noWrap")
        tc_pr.append(no_wrap)


def _arcname_for(run_root: Path, path: Path) -> str:
    try:
        return path.resolve().relative_to(run_root.resolve()).as_posix()
    except ValueError:
        parts = list(path.resolve().parts)
        if run_root.name in parts:
            index = len(parts) - 1 - parts[::-1].index(run_root.name)
            tail = parts[index + 1:]
            if tail:
                return Path(*tail).as_posix()
        return safe_filename(path.name)


def _unique_arcname(arcname: str, seen: set[str]) -> str:
    candidate = arcname
    if candidate not in seen:
        seen.add(candidate)
        return candidate
    path = Path(arcname)
    index = 2
    while True:
        candidate = str(path.with_name(f"{path.stem}_{index}{path.suffix}")).replace("\\", "/")
        if candidate not in seen:
            seen.add(candidate)
            return candidate
        index += 1


def _missing_detail(row: ResultRow, path: Path) -> dict:
    return {
        "task_name": row.task_name,
        "device_group": row.device_group,
        "device_name": row.device_name,
        "path": str(path),
    }


def _norm(value: str) -> str:
    return " ".join(str(value or "").split())
